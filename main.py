from flask import Flask, request, jsonify, send_file
import openai
import zipfile
import io
import os
from fpdf import FPDF
from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path
from dotenv import load_dotenv

app = Flask(__name__)

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")

# Function to fetch GPT response
def fetch_gpt_response(query):
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an assistant answering only questions related to the given PDF content. Don't answer any other irrelevant questions."},
                {"role": "user", "content": query},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

# Function to extract text from a PDF
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    extracted_text = ""
    for page in reader.pages:
        extracted_text += page.extract_text()
    return extracted_text

# Function to save SCORM package with PDF
def save_as_scorm_pdf(content, output_folder="scorm_package", scorm_zip_name="scorm_pdf_package.zip"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    pdf_path = os.path.join(output_folder, "response.pdf")
    save_as_pdf(content, pdf_path)

    html_path = os.path.join(output_folder, "index.html")
    with open(html_path, "w") as html_file:
        html_file.write(f"""
        <!DOCTYPE html>
        <html>
        <head><title>PDF Analysis</title></head>
        <body>
        <h1>PDF Analysis Response</h1>
        <iframe src="response.pdf" width="100%" height="600px"></iframe>
        </body>
        </html>
        """)

    manifest_path = os.path.join(output_folder, "imsmanifest.xml")
    with open(manifest_path, "w") as manifest_file:
        manifest_file.write(f"""
        <?xml version="1.0" encoding="UTF-8"?>
        <manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
                  xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_v1p3"
                  xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1">
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <organizations>
                <organization identifier="ORG-1">
                    <title>PDF Analysis Response</title>
                </organization>
            </organizations>
            <resources>
                <resource identifier="RES-1" type="webcontent" href="index.html">
                    <file href="index.html"/>
                    <file href="response.pdf"/>
                </resource>
            </resources>
        </manifest>
        """)

    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        for root, _, files in os.walk(output_folder):
            for file in files:
                scorm_zip.write(os.path.join(root, file), arcname=file)

    return scorm_zip_name

# Function to save SCORM package with Word document
def save_as_scorm_word(content, output_folder="scorm_package", scorm_zip_name="scorm_doc_package.zip"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    doc_path = os.path.join(output_folder, "response.docx")
    doc = Document()
    doc.add_paragraph(content)
    doc.save(doc_path)

    html_path = os.path.join(output_folder, "index.html")
    with open(html_path, "w") as html_file:
        html_file.write(f"""
        <!DOCTYPE html>
        <html>
        <head><title>PDF Analysis</title></head>
        <body>
        <h1>PDF Analysis Response</h1>
        <p>{content.replace('\n', '<br>')}</p>
        </body>
        </html>
        """.format(content.replace('\n', '<br>'))

    manifest_path = os.path.join(output_folder, "imsmanifest.xml")
    with open(manifest_path, "w") as manifest_file:
        manifest_file.write(f"""
        <?xml version="1.0" encoding="UTF-8"?>
        <manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
                  xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_v1p3"
                  xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1">
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <organizations>
                <organization identifier="ORG-1">
                    <title>PDF Analysis Response</title>
                </organization>
            </organizations>
            <resources>
                <resource identifier="RES-1" type="webcontent" href="index.html">
                    <file href="index.html"/>
                    <file href="response.docx"/>
                </resource>
            </resources>
        </manifest>
        """)

    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        for root, _, files in os.walk(output_folder):
            for file in files:
                scorm_zip.write(os.path.join(root, file), arcname=file)

    return scorm_zip_name

def save_as_pdf(content, file_name="response.pdf", logo_path='assets/logo.jpeg'):
    pdf = FPDF()
    pdf.add_page()

    # Add the logo
    pdf.image(logo_path, x=10, y=8, w=30)

    # Title of the document
    pdf.set_font("Arial", style='B', size=16)
    pdf.ln(30)
    pdf.cell(200, 10, txt="Research Content Response", ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, content)

    # Save the PDF
    pdf.output(file_name)

def save_as_word(content, file_name="response.docx", logo_path='assets/logo.jpeg'):
    document = Document()

    # Add the logo
    document.add_picture(logo_path, width=1500000)  # Width in EMU (e.g., 1500000 = 150px)

    # Add the title
    document.add_heading("Research Content Response", level=1)

    # Add content
    document.add_paragraph(content)

    # Save the Word document
    document.save(file_name)

@app.route('/api/analyze-pdf', methods=['POST'])
def analyze_pdf():
    try:
        if 'pdf_file' not in request.files:
            return jsonify({"error": "No PDF uploaded."}), 400

        pdf_file = request.files['pdf_file']
        pdf_text = extract_text_from_pdf(pdf_file)

        query = request.form.get('query', '').strip()
        if not query:
            return jsonify({"error": "Query is required."}), 400

        response = fetch_gpt_response(f"Context: {pdf_text}\nQuestion: {query}")

        return jsonify({"pdf_text": pdf_text, "response": response})

    except Exception as e:
        return jsonify({"error": f"Error: {str(e)}"}), 500

@app.route('/api/download-scorm', methods=['POST'])
def download_scorm():
    try:
        content = request.json.get('response', '')
        scorm_type = request.json.get('scorm_type', '')

        if scorm_type == 'pdf':
            scorm_zip_path = save_as_scorm_pdf(content)
        elif scorm_type == 'doc':
            scorm_zip_path = save_as_scorm_word(content)
        else:
            return jsonify({"error": "Invalid SCORM type selected."}), 400

        return send_file(
            scorm_zip_path,
            as_attachment=True,
            download_name="scorm_package.zip",
            mimetype="application/zip"
        )

    except Exception as e:
        return jsonify({"error": f"Error: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True)
